import os, tempfile
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def add_horizontal_line(paragraph):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '6C63FF')
    pBdr.append(bottom)
    pPr.append(pBdr)


def create_docx(data: dict) -> str:
    doc = Document()
    info = data.get("personal", {})
    text = data.get("resumeText", "")

    for section in doc.sections:
        section.top_margin    = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin   = Inches(0.85)
        section.right_margin  = Inches(0.85)

    # Name
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = name_para.add_run(info.get("name", "Your Name"))
    run.font.size  = Pt(22)
    run.font.bold  = True
    run.font.color.rgb = RGBColor(0x18, 0x18, 0x1B)

    # Contact
    contact_parts = [x for x in [
        info.get("email"), info.get("phone"), info.get("location"),
        info.get("linkedin"), info.get("github"), info.get("website"),
    ] if x]
    cp = doc.add_paragraph("  |  ".join(contact_parts))
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cp.runs:
        run.font.size  = Pt(9)
        run.font.color.rgb = RGBColor(0x71, 0x71, 0x7A)
    add_horizontal_line(cp)
    doc.add_paragraph()

    # Body
    for line in text.split("\n"):
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue
        if line.isupper() or (line.endswith(":") and len(line) < 40):
            h = doc.add_heading("", level=2)
            run = h.add_run(line.rstrip(":"))
            run.font.color.rgb = RGBColor(0x6C, 0x63, 0xFF)
            run.font.size = Pt(11)
        elif line.startswith(("•", "-", "*")):
            p = doc.add_paragraph(line.lstrip("•-* "), style="List Bullet")
            for r in p.runs:
                r.font.size = Pt(10)
        else:
            p = doc.add_paragraph(line)
            for r in p.runs:
                r.font.size = Pt(10)

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(tmp.name)
    return tmp.name
